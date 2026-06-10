"""
취업규칙 docx → ChromaDB 임베딩 빌드 스크립트.

실행 (chatbot/ 디렉터리에서):
    python scripts/build_company_rules_db.py

한 번만 실행하면 됩니다. 이미 collection이 있으면 재빌드 여부를 묻습니다.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).parent.parent          # chatbot/
sys.path.insert(0, str(ROOT))

from docx import Document
from langchain.schema import Document as LCDocument
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma

import config

DOCX_PATH = ROOT / "data" / "company_rules" / "raw" / "company_work_rules_masked_v1.docx"
COLLECTION_NAME = "company_rules"

ARTICLE_RE = re.compile(r"^(제\d+조(?:의\d+)?)\s*【([^】]+)】\s*$")
CHAPTER_RE = re.compile(r"^(제\d+장)\s+(.+?)\s*$")


def _find_body_start(paragraphs: list) -> int:
    """TOC 끝, 본문 시작 위치 탐색 (tab 없는 첫 번째 제1조 위치)."""
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if re.match(r"^제1조\s*【목\s*적】\s*$", text):
            return i
    return 180  # fallback


def parse_articles(docx_path: Path) -> list[LCDocument]:
    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs

    body_start = _find_body_start(paragraphs)
    print(f"[parse] 본문 시작: 단락 {body_start}")

    docs: list[LCDocument] = []
    chapter_no = ""
    chapter_title = ""
    article_no = ""
    article_title = ""
    content_lines: list[str] = []

    def flush():
        nonlocal article_no, article_title, content_lines
        if article_no and content_lines:
            body = "\n".join(content_lines).strip()
            if body:
                page_content = f"[{article_no} {article_title}]\n{body}"
                docs.append(LCDocument(
                    page_content=page_content,
                    metadata={
                        "chapter_no": chapter_no,
                        "chapter_title": chapter_title,
                        "article_no": article_no,
                        "article_title": article_title,
                        "source": "취업규칙",
                        "company": "맘마중컴퍼니 코리아",
                        "type": "사내규정",
                    },
                ))
        article_no = ""
        article_title = ""
        content_lines = []

    for p in paragraphs[body_start:]:
        text = p.text.strip()
        if not text:
            continue

        chapter_m = CHAPTER_RE.match(text)
        article_m = ARTICLE_RE.match(text)

        if chapter_m:
            flush()
            chapter_no = chapter_m.group(1)
            chapter_title = chapter_m.group(2).strip()
        elif article_m:
            flush()
            article_no = article_m.group(1)
            article_title = article_m.group(2).strip()
        else:
            if article_no:
                content_lines.append(text)

    flush()
    return docs


def build(force: bool = False) -> None:
    embeddings = OpenAIEmbeddings(
        model="text-embedding-3-small",
        openai_api_key=config.OPENAI_API_KEY,
    )

    # 기존 컬렉션 확인
    import chromadb
    client = chromadb.PersistentClient(path=config.CHROMA_PATH)
    existing = client.list_collections()  # v0.6+: returns list of names (str)
    if COLLECTION_NAME in existing:
        if not force:
            ans = input(f"[빌드] '{COLLECTION_NAME}' 컬렉션이 이미 있습니다. 재빌드하겠습니까? (y/N) ").strip().lower()
            if ans != "y":
                print("[빌드] 취소.")
                return
        client.delete_collection(COLLECTION_NAME)
        print(f"[빌드] 기존 '{COLLECTION_NAME}' 삭제 완료.")

    print(f"[빌드] {DOCX_PATH.name} 파싱 중…")
    articles = parse_articles(DOCX_PATH)
    print(f"[빌드] 파싱 완료: {len(articles)}개 조문")

    if not articles:
        print("[빌드] 오류: 파싱된 조문이 없습니다.")
        sys.exit(1)

    print("[빌드] ChromaDB 임베딩 중… (OpenAI API 호출)")
    vectorstore = Chroma.from_documents(
        articles,
        embeddings,
        persist_directory=config.CHROMA_PATH,
        collection_name=COLLECTION_NAME,
    )
    print(f"[빌드] 완료: '{COLLECTION_NAME}' 컬렉션에 {len(articles)}개 조문 저장.")

    # 간단 검색 테스트
    test_results = vectorstore.similarity_search("배우자 출산휴가 기간", k=2)
    print("\n[테스트] '배우자 출산휴가 기간' 검색 결과:")
    for r in test_results:
        print(f"  → {r.metadata['article_no']} {r.metadata['article_title']}: {r.page_content[:60]}…")


if __name__ == "__main__":
    force_flag = "--force" in sys.argv
    build(force=force_flag)
